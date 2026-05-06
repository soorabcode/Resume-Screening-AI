"""
resume_screening_ai.py  —  Complete Resume Screening AI in one file
=====================================================================
Usage (CLI):
    python resume_screening_ai.py --resumes ./resumes/ --jd job.txt
    python resume_screening_ai.py --resumes ./resumes/ --jd job.txt --min-score 50 --output results.csv

Usage (Jupyter/Colab):
    # Run the cell once to import functions
    from resume_screening_ai import run_pipeline, extract_all, parse_resume
    # Construct sources manually and call:
    # results = run_pipeline(sources, jd_text, min_score=50.0, output_path="results.csv")
"""

import re, io, argparse, sys
from pathlib import Path
from datetime import datetime

import numpy as np
import pandas as pd
from dateutil import parser as dateparser
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer

# ══════════════════════════════════════════════════════════════════════════════
# 1. RESUME PARSER  —  extract raw text from PDF / DOCX
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(source) -> str:
    if isinstance(source, (str, Path)):
        with open(source, "rb") as f:
            data = f.read()
    elif isinstance(source, bytes):
        data = source
    else:
        data = source.read()

    text = ""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype="pdf")
        for page in doc:
            text += page.get_text()
        doc.close()
        if text.strip():
            return text
    except Exception:
        pass

    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception as e:
        print(f"[PDF] Error: {e}")
    return text


def extract_text_from_docx(source) -> str:
    from docx import Document
    if isinstance(source, (str, Path)):
        doc = Document(source)
    else:
        doc = Document(io.BytesIO(source) if isinstance(source, bytes) else source)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def clean_text(text: str) -> str:
    text = re.sub(r"[^\S\n]+", " ", text)
    text = re.sub(r"[^\x20-\x7E\n]", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_resume(source, filename: str = "") -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return clean_text(extract_text_from_pdf(source))
    elif name.endswith(".docx"):
        return clean_text(extract_text_from_docx(source))
    raise ValueError(f"Unsupported file type: {filename}")


# ══════════════════════════════════════════════════════════════════════════════
# 2. INFORMATION EXTRACTOR  —  skills, education, experience, name
# ══════════════════════════════════════════════════════════════════════════════

SKILLS_MASTER = [
    "python","java","javascript","typescript","c++","c#","go","rust","ruby","php",
    "swift","kotlin","scala","r","matlab","html","css","react","angular","vue",
    "nextjs","nodejs","express","django","flask","fastapi","spring","rails",
    "machine learning","deep learning","nlp","computer vision","tensorflow","pytorch",
    "keras","scikit-learn","xgboost","lightgbm","hugging face","transformers","bert",
    "gpt","llm","rag","pandas","numpy","scipy","matplotlib","seaborn","plotly",
    "sql","mysql","postgresql","mongodb","redis","elasticsearch","spark","hadoop",
    "airflow","kafka","dbt","snowflake","bigquery","aws","azure","gcp","docker",
    "kubernetes","terraform","ansible","ci/cd","jenkins","github actions","linux",
    "bash","git","agile","scrum","rest api","graphql","microservices","statistics",
    "a/b testing","excel","tableau","power bi","looker","data analysis",
    "data visualization","faiss","pdfplumber","sentence-transformers",
]

DEGREE_KEYWORDS = [
    "bachelor","master","phd","doctorate","b.sc","m.sc","b.e","m.e",
    "b.tech","m.tech","mba","b.a","m.a","associate","diploma",
    "undergraduate","postgraduate",
]

SKIP_WORDS = {
    "resume","curriculum","vitae","cv","profile","summary",
    "objective","contact","email","phone","address",
}


def extract_name(text: str) -> str:
    for line in [l.strip() for l in text.split("\n") if l.strip()][:6]:
        words = line.split()
        if 2 <= len(words) <= 4:
            if all(w[0].isupper() and w.isalpha() for w in words):
                if not any(w.lower() in SKIP_WORDS for w in words):
                    return line
    return "Unknown"


def extract_email(text: str):
    m = re.search(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}", text)
    return m.group(0) if m else None


def extract_skills(text: str, vocab: list = None) -> list:
    vocab = vocab or SKILLS_MASTER
    text_lower = text.lower()
    found = set()
    for s in vocab:
        # Non-word boundary assertions correctly handle punctuation in skills (e.g., C++, C#)
        pattern = r"(?<!\w)" + re.escape(s) + r"(?!\w)"
        if re.search(pattern, text_lower):
            found.add(s)
    return sorted(found)


def extract_education(text: str) -> list:
    results = []
    for line in text.split("\n"):
        if any(kw in line.lower() for kw in DEGREE_KEYWORDS):
            if line.strip() and len(line.strip()) > 5:
                results.append(line.strip())
    return results[:5]


def extract_years_of_experience(text: str) -> float:
    total_months = 0.0
    pattern = re.compile(
        r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})"
        r"\s*[-–—to]+\s*"
        r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|[Pp]resent|[Cc]urrent|[Nn]ow)",
        re.IGNORECASE,
    )
    now = datetime.now()
    for m in pattern.finditer(text):
        try:
            start = dateparser.parse(m.group(1), default=datetime(now.year, 1, 1))
            end_s = m.group(2)
            end = now if re.match(r"present|current|now", end_s, re.I) else dateparser.parse(end_s, default=datetime(now.year, 1, 1))
            if start and end and end >= start:
                total_months += (end.year - start.year) * 12 + (end.month - start.month)
        except Exception:
            continue

    if total_months > 0:
        return round(total_months / 12, 1)

    m = re.search(r"(\d+)\+?\s+years?\s+(?:of\s+)?experience", text, re.I)
    return float(m.group(1)) if m else 0.0


def extract_required_years(jd_text: str) -> float:
    m = re.compile(r"(\d+)\s*\+?\s*(?:to\s*\d+\s*)?years?\s+(?:of\s+)?(?:relevant\s+)?experience", re.I).search(jd_text)
    if m:
        return float(m.group(1))
    m2 = re.search(r"(\d+)\s*[-–]\s*\d+\s+years?", jd_text, re.I)
    return float(m2.group(1)) if m2 else 0.0


def extract_all(text: str) -> dict:
    return {
        "name": extract_name(text),
        "email": extract_email(text),
        "skills": extract_skills(text),
        "education": extract_education(text),
        "years_experience": extract_years_of_experience(text),
        "raw_text": text,
    }


# ══════════════════════════════════════════════════════════════════════════════
# 3. SCORER  —  embeddings + weighted sub-scores
# ══════════════════════════════════════════════════════════════════════════════

_MODEL = None

def get_model(name: str = "all-MiniLM-L6-v2") -> SentenceTransformer:
    global _MODEL
    if _MODEL is None:
        print(f"[Model] Loading {name} …")
        _MODEL = SentenceTransformer(name)
    return _MODEL


def embed(texts: list) -> np.ndarray:
    return get_model().encode(texts, convert_to_numpy=True, show_progress_bar=False)


def score_skills(resume_skills: list, jd_skills: list) -> float:
    if not jd_skills:
        return 0.5
    rs, js = set(resume_skills), set(jd_skills)
    return len(rs & js) / len(js)


def score_experience(candidate_yrs: float, required_yrs: float) -> float:
    if required_yrs <= 0:
        return 0.6
    return float(np.clip(candidate_yrs / required_yrs, 0.0, 1.0))


def score_semantic(resume_text: str, jd_text: str) -> float:
    vecs = embed([resume_text, jd_text])
    sim = cosine_similarity(vecs[0:1], vecs[1:2])[0][0]
    return float(np.clip(sim, 0.0, 1.0))


def compute_score(resume_text, resume_skills, jd_text, jd_skills,
                  candidate_yrs, required_yrs,
                  weights=None) -> dict:
    w = weights or {"skills": 0.40, "experience": 0.30, "semantic": 0.30}
    ss = score_skills(resume_skills, jd_skills)
    se = score_experience(candidate_yrs, required_yrs)
    sm = score_semantic(resume_text, jd_text)
    total = w["skills"] * ss + w["experience"] * se + w["semantic"] * sm
    return {
        "skills_score":     round(ss * 100, 1),
        "experience_score": round(se * 100, 1),
        "semantic_score":   round(sm * 100, 1),
        "total_score":      round(total * 100, 1),
    }


# ══════════════════════════════════════════════════════════════════════════════
# 4. EXPLAINER  —  readable summaries + keyword highlighting
# ══════════════════════════════════════════════════════════════════════════════

def skill_diff(resume_skills: list, jd_skills: list) -> dict:
    rs, js = set(resume_skills), set(jd_skills)
    return {"matched": sorted(rs & js), "missing": sorted(js - rs), "extra": sorted(rs - js)}


def explain(name, score_dict, diff, candidate_yrs, required_yrs, rank) -> str:
    ordinal = {1:"1st",2:"2nd",3:"3rd"}.get(rank, f"{rank}th")
    lines = [f"**{name}** ranked {ordinal} with {score_dict['total_score']:.1f}/100."]

    n_m, n_x = len(diff["matched"]), len(diff["missing"])
    if n_m == 0:
        lines.append("No required skills found in the resume.")
    elif n_x == 0:
        lines.append(f"All {n_m} required skill(s) matched — perfect skills fit.")
    else:
        lines.append(f"{n_m} skill(s) matched ({', '.join(diff['matched'][:5])}); "
                     f"{n_x} missing ({', '.join(diff['missing'][:5])}).")

    if required_yrs > 0:
        if candidate_yrs >= required_yrs:
            lines.append(f"Experience ({candidate_yrs:.1f} yrs) meets the {required_yrs:.0f}+ yr requirement.")
        else:
            lines.append(f"Experience ({candidate_yrs:.1f} yrs) is {required_yrs-candidate_yrs:.1f} yr(s) short.")
    else:
        lines.append(f"No explicit experience requirement; candidate has ~{candidate_yrs:.1f} yr(s) on record.")

    sem = score_dict["semantic_score"]
    tag = "very high" if sem >= 75 else "moderate" if sem >= 50 else "low"
    lines.append(f"Semantic alignment: {tag} ({sem:.0f}/100).")
    return " ".join(lines)


def highlight_keywords(text: str, keywords: list) -> str:
    snippet = text[:2000]
    for kw in sorted(keywords, key=len, reverse=True):
        snippet = re.compile(r"\b(" + re.escape(kw) + r")\b", re.I).sub(r"**\1**", snippet)
    return snippet


# ══════════════════════════════════════════════════════════════════════════════
# 5. PIPELINE  —  orchestrates everything
# ══════════════════════════════════════════════════════════════════════════════

def run_pipeline(resume_sources: list, jd_text: str,
                 min_score: float = 0.0, output_path: str = None,
                 weights: dict = None) -> list:
    """
    resume_sources: [{"bytes": b"...", "filename": "cv.pdf"}, ...]
    Returns ranked list of result dicts.
    """
    SEP = "═" * 60
    print(f"\n{SEP}\n  RESUME SCREENING AI  –  Pipeline Starting\n{SEP}")

    # Step 1: JD
    print("\n[1/4] Processing job description …")
    jd_skills = extract_skills(jd_text)
    required_yrs = extract_required_years(jd_text)
    print(f"      Skills detected : {len(jd_skills)}  |  Required exp: {required_yrs} yrs")

    # Step 2: Parse resumes
    print(f"\n[2/4] Parsing {len(resume_sources)} resume(s) …")
    parsed = []
    for src in resume_sources:
        try:
            info = extract_all(parse_resume(src["bytes"], src["filename"]))
            info["filename"] = src["filename"]
            parsed.append(info)
            print(f"      ✓ {src['filename']}  →  {info['name']}")
        except Exception as e:
            print(f"      ✗ {src['filename']}  →  ERROR: {e}")

    if not parsed:
        print("[ERROR] No resumes parsed successfully.")
        return []

    # Step 3: Score
    print("\n[3/4] Scoring candidates …")
    for info in parsed:
        scores = compute_score(
            resume_text=info["raw_text"], resume_skills=info["skills"],
            jd_text=jd_text, jd_skills=jd_skills,
            candidate_yrs=info["years_experience"], required_yrs=required_yrs,
            weights=weights,
        )
        info.update(scores)

    # Step 4: Rank + explain
    print("\n[4/4] Ranking …")
    parsed.sort(key=lambda x: x["total_score"], reverse=True)

    ranked = []
    for rank, r in enumerate(parsed, 1):
        if r["total_score"] < min_score:
            continue
        diff = skill_diff(r["skills"], jd_skills)
        ranked.append({
            "rank":             rank,
            "name":             r["name"],
            "email":            r.get("email", ""),
            "file":             r["filename"],
            "total_score":      r["total_score"],
            "skills_score":     r["skills_score"],
            "experience_score": r["experience_score"],
            "semantic_score":   r["semantic_score"],
            "years_experience": r["years_experience"],
            "matched_skills":   ", ".join(diff["matched"]),
            "missing_skills":   ", ".join(diff["missing"]),
            "education":        " | ".join(r.get("education", [])),
            "explanation":      explain(r["name"], r, diff, r["years_experience"], required_yrs, rank),
            "highlighted_snippet": highlight_keywords(r["raw_text"], diff["matched"]),
        })

    # Print results
    print(f"\n{SEP}\n  RESULTS\n{SEP}")
    df = pd.DataFrame(ranked)
    if not df.empty:
        cols = ["rank","name","total_score","skills_score","experience_score","semantic_score","years_experience"]
        print(df[cols].to_string(index=False))
        if output_path:
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            df.to_csv(output_path, index=False)
            print(f"\n[✓] Saved → {output_path}")
    else:
        print("  No candidates met the minimum score threshold.")
    print(SEP + "\n")
    return ranked


# ══════════════════════════════════════════════════════════════════════════════
# 6. CLI & INTERACTIVE COMPATIBLE ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

def main():
    ap = argparse.ArgumentParser(description="Resume Screening AI")
    ap.add_argument("--resumes",   required=True, help="Folder with PDF/DOCX resumes")
    ap.add_argument("--jd",        required=True, help="Job description .txt file")
    ap.add_argument("--min-score", type=float, default=0.0, help="Min score 0-100")
    ap.add_argument("--output",    default="results.csv", help="Output CSV path")

    try:
        args = ap.parse_args()
    except SystemExit:
        # Gracefully handle missing arguments in interactive environments (Jupyter/Colab)
        print("\n[INFO] Command-line arguments are required for terminal execution.")
        print("CLI Usage: python resume_screening_ai.py --resumes ./resumes/ --jd job.txt")
        print("\nFor Jupyter/Colab, please execute the pipeline directly:")
        print("  sources = [{'bytes': open(f,'rb').read(), 'filename': f.name} for f in Path('./resumes/').iterdir()]")
        print("  jd_text = Path('job.txt').read_text()")
        print("  run_pipeline(sources, jd_text, min_score=50.0)")
        return

    jd_path = Path(args.jd)
    if not jd_path.exists():
        sys.exit(f"[ERROR] JD not found: {jd_path}")
    jd_text = jd_path.read_text(encoding="utf-8", errors="ignore")

    resume_dir = Path(args.resumes)
    if not resume_dir.is_dir():
        sys.exit(f"[ERROR] Directory not found: {resume_dir}")

    sources = [
        {"bytes": fp.read_bytes(), "filename": fp.name}
        for fp in sorted(resume_dir.iterdir())
        if fp.suffix.lower() in (".pdf", ".docx")
    ]

    if not sources:
        sys.exit("[ERROR] No PDF/DOCX files found.")

    run_pipeline(sources, jd_text, min_score=args.min_score, output_path=args.output)


if __name__ == "__main__":
    main()